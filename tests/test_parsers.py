import pytest
import os
from parsers import parse_file

FIXTURES = os.path.join(os.path.dirname(__file__), "fixtures")


def test_parse_txt():
    path = os.path.join(FIXTURES, "sample.txt")
    result = parse_file(path)
    assert result["text"].strip() != ""
    assert result["filename"] == "sample.txt"


def test_parse_markdown():
    path = os.path.join(FIXTURES, "sample.md")
    result = parse_file(path)
    assert len(result["text"]) > 0
    # Should strip markdown formatting
    assert "**" not in result["text"]


def test_parse_html():
    path = os.path.join(FIXTURES, "sample.html")
    result = parse_file(path)
    assert "<" not in result["text"]  # Tags stripped
    assert len(result["text"]) > 0
    assert "Sample Page" in result["text"]
    # Script and style content should be removed
    assert "var x" not in result["text"]
    assert "display" not in result["text"]


def test_parse_unsupported():
    with pytest.raises(ValueError, match="Unsupported"):
        parse_file("test.xyz")


def test_parse_returns_pages_for_pdf(tmp_path):
    """PDF parser should return page_numbers metadata."""
    import fitz  # PyMuPDF

    pdf_path = tmp_path / "test.pdf"
    doc = fitz.open()
    page = doc.new_page()
    writer = fitz.TextWriter(page.rect)
    writer.append((72, 72), "Page one content here", fontsize=12)
    writer.write_text(page)
    doc.save(str(pdf_path))
    doc.close()

    result = parse_file(str(pdf_path))
    assert "Page one content" in result["text"]
    assert result.get("pages") is not None


def test_parse_docx(tmp_path):
    """DOCX parser extracts paragraph text."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello from docx")
    doc.add_paragraph("Second paragraph")
    path = tmp_path / "test.docx"
    doc.save(str(path))

    result = parse_file(str(path))
    assert "Hello from docx" in result["text"]
    assert "Second paragraph" in result["text"]
